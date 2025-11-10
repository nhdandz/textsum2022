import base64
import tempfile
import sys, fitz
import docx
import re
import os
import aspose.words as aw
TEXT = 0
PDF = 1
DOCX = 2
DOC =3

def get_raw_text(raw_text,file_type,page_from,page_to):
    if file_type == TEXT:
        return raw_text
    elif file_type == PDF or file_type == DOC or file_type == DOCX:
        ischeck =  isBase64(raw_text)
        if ischeck == True:
            if page_to < 0:
                page_to = 0
            if file_type == PDF:
                temp = tempfile.NamedTemporaryFile(suffix='.pdf', delete=True)
                print("Created file is:", temp)
                print("Name of the file is:", temp.name)
                decode_raw =  base64.b64decode(raw_text)
                binary_file = open(temp.name, "wb")
                # Write bytes to file
                binary_file.write(decode_raw)
                text = process_pdf_text(temp.name,page_from,page_to)
            else:
                temp = tempfile.NamedTemporaryFile(suffix='.docx', delete=True)
                print("Created file is:", temp)
                print("Name of the file is:", temp.name)
                decode_raw =  base64.b64decode(raw_text)
                binary_file = open(temp.name, "wb")
                # Write bytes to file
                binary_file.write(decode_raw)
                text = process_doc_text(temp.name,page_from,page_to)
            return text
        else:
            return ''
    else:
        return ''
def get_number_page(raw_text,file_type):
    if file_type == TEXT:
        return raw_text
    elif file_type == PDF or file_type == DOC or file_type == DOCX:
        ischeck =  isBase64(raw_text)
        if ischeck == True:
            temp = tempfile.NamedTemporaryFile()
            print("Created file is:", temp)
            print("Name of the file is:", temp.name)
            decode_raw =  base64.b64decode(raw_text)
            binary_file = open(temp.name, "wb")
            # Write bytes to file
            binary_file.write(decode_raw)
            
            if file_type == PDF:
                doc = fitz.open(temp.name) 
                page_final = doc.page_count
            else:
                doc = aw.Document(temp.name)
                page = doc.page_count + 1
                # xet truong hop dac biet page=2
                # khi nay text_page_2 ==text_page_1 do luon tinh tu trang 
                if page > 2:
                    text_page_2 = process_doc_text(temp.name, page-2, page-2)
                    text_page_1 = process_doc_text(temp.name, page-1, page-1)
                    if text_page_1 == '' and text_page_2 =='':
                        page_final = page - 2
                    elif text_page_1 =='' and text_page_2 !='':
                        page_final = page - 1
                    else:
                        page_final = page
                else:
                    page_final = doc.page_count
            return page_final
        else:
            return 0
    else:
        return 0
def delete_if_exist(filePath):
    if os.path.exists(filePath):
        os.remove(filePath)
    else:
        print("Can not delete the file as it doesn't exists")
def get_raw_text_by_topic(topics,raw_text):
    list_pragrab = raw_text.split('\n')
    list_raw = []
    topic_choose = topics[0]
    topic_not = topics[1]
    if len(topic_choose) != 0:
        for key_words in topic_choose:
            key_word = key_words.split(',')
            for pragrab in list_pragrab :
                is_choose = True
                for word in key_word:
                    if word.lower() not in pragrab.lower():
                        is_choose = False
                if is_choose == True:
                    index = list_pragrab.index(pragrab)
                    list_raw.append((index,pragrab))
    else:
        count = 0
        for pragrab in list_pragrab :
            list_raw.append((count,pragrab))
            count +=1
    list_raw = removeDuplicates(list_raw)
    list_raw_process = sorted(list_raw, key=lambda tup: tup[0])
    list_raw_final = []
    if len(topic_not) != 0:
        for key_words in topic_not:
            key_word = key_words.split(',')
            for pragrab in list_raw_process :
                is_choose = True
                for word in key_word:
                    if ' '+word.strip().lower()+' ' in pragrab[1].lower():
                        is_choose = False
                if is_choose == True:
                    list_raw_final.append(pragrab)
    else:
        list_raw_final = list_raw_process
    list_raw_final = removeDuplicates(list_raw_final)
    list_raw_final_process = sorted(list_raw_final, key=lambda tup: tup[0])
    list_text_topic = []
    for text_topic in list_raw_final_process:
        list_text_topic.append(text_topic[1])
    return '\n'.join(list_text_topic)
def check_id_mapAlgTypeAI(mapAlgTypeAI,id_mapAlgTypeAI):
    try:
        if len (id_mapAlgTypeAI) == 2:
            for item in mapAlgTypeAI:
                if item['id'] == id_mapAlgTypeAI[0]:
                    typeAIId_short =  item['typeAIId']
                    aicore_short = item['aiId']
            for item in mapAlgTypeAI:
                if item['id'] == id_mapAlgTypeAI[1]:
                    typeAIId_long =  item['typeAIId']
                    aicore_long = item['aiId']
            if typeAIId_short == typeAIId_long and aicore_short == 1 and aicore_long == 2:
                return True
        else: 
            return False
    except:
        return False
def check_percent_output(percent_output):
    try:
        if (isinstance(percent_output, float) or percent_output==0)  and percent_output >=0 and percent_output <= 1:
            return True
        else: 
            return False
    except:
        return False
def check_topic(topic):
    if len(topic) == 2:
        return True
    else:
        return False
def check_add_optional_value(content):
    content.setdefault("topic", {})
    content.setdefault("percent_output", 0.3)
    content.setdefault("page_from", 0)
    content.setdefault("page_to", 99999)

def check_add_optional_value_document(content):
    content.setdefault("page_from", 0)
    content.setdefault("page_to", 99999)

def check_add_optional_value_msg(content):
    content.setdefault("percent_output", 0.3)
    content.setdefault("topic", [])

def removeDuplicates(lst):
    return list(set([i for i in lst]))
def get_length(raw_text):
    paragraph_list = raw_text.split('\n')
    word =  ''.join(paragraph_list)
    word_list = word.split(' ')
    number_of_words = len(word_list)
    return number_of_words
def check_short(text):
    len_text = get_length(text)
    if len_text > 800:
        return False
    else:
        return True
def isBase64(sb):
    try:
        if isinstance(sb, str):
            sb_bytes = sb.encode('utf-8')  # ✅ đổi thành utf-8
        elif isinstance(sb, bytes):
            sb_bytes = sb
        else:
            return False
        return base64.b64encode(base64.b64decode(sb_bytes)) == sb_bytes
    except Exception:
        return False
     
# def process_pdf_text(file_name,page_from,page_to):
#     doc = fitz.open(file_name)  # open document
#     all_text = ''
#     for page in doc:  # iterate the document pages
#         if page.number >= page_from and page.number <= page_to:
#             blocks = page.get_text("blocks")  # get plain text (is in UTF-8)
#             # process into paragraph
#             list_paragraph = []
#             for b in blocks:
#                 temp_paragraph = re.split('\n\n| \n \n',b[4])
#                 for paragraph in temp_paragraph:
#                     list_paragraph.append(paragraph)
#             list_paragraph_process = []
#             for paragraph in list_paragraph:
#                 text = paragraph.replace("\n", " ")
#                 re.sub('[^A-Za-z0-9]+', '', text)
#                 text = re.sub(' +', ' ', text)
#                 strencode = text.encode("ascii", "ignore")
#                 #decode() method
#                 text = strencode.decode()
#                 text = text.rstrip()
#                 text = text.strip()
#                 if '<image' in text:
#                     continue
#                 list_paragraph_process.append(text)
#             for idx, text in enumerate(list_paragraph_process):
#                 if text =='':
#                     continue
#                 if get_length(text) > 25 and text[0].isupper() == True :
#                     all_text  = all_text+'\n'+ text
#                     continue
#                 if text[0].isupper() == True and text[-1] != '.':
#                     try:
#                         if len(list_paragraph_process[idx+1])!=0 and list_paragraph_process[idx+1][0].isupper() == False and list_paragraph_process[idx+1][-1] =='.':
#                             list_paragraph_process[idx] = list_paragraph_process[idx]+list_paragraph_process[idx+1]
#                             all_text  = all_text+'\n'+ list_paragraph_process[idx]
#                             list_paragraph_process.pop(idx+1)
#                     except:
#                         a =5
#     return all_text

def process_pdf_text(file_name,page_from,page_to):
    doc = fitz.open(file_name)  # open document
    all_text = ''
    count = 0
    all_text_list_process = []
    all_text_list_unprocess = []
    stop_paragraph = ('...', '.', '?', '!', '!!!', ':')
    for page in doc:  # iterate the document pages
        if page.number >= page_from and page.number <= page_to:
            count+=1
            blocks = page.get_text("blocks")  # get plain text (is in UTF-8)
            # process into paragraph
            list_paragraph = []
            for b in blocks:
                temp_paragraph = re.split('\n\n| \n \n|\n \n',re.sub(' +', ' ', b[4]))
                for paragraph in temp_paragraph:
                    list_paragraph.append(paragraph)
            all_text_list = []
            for paragraph in list_paragraph:
                text = paragraph.replace("\n", " ") 
                # strencode = text.encode("ascii", "ignore")
                strencode = text.encode()
                #decode() method
                text = strencode.decode()
                text = text.replace("\t", "") 
                text = text.replace("\r", "") 
                # re.sub('[^A-Za-z0-9]+', '', text)
                re.sub(' +', ' ', text)
                text = text.rstrip()
                text = text.strip()
                if '<image' in text:
                    continue
                if isinstance(re.search("Figure [0-9]", text), tuple) and re.search("Figure [0-9]", text).span()[0] == 0:
                    continue
                if isinstance(re.search("Table [0-9]", text), tuple) and re.search("Table [0-9]", text).span()[0] == 0:
                    continue
                if 'VerDate' in text:
                    continue
                all_text_list.append(text)
                if text != '':
                    all_text_list_unprocess.append(text)
            for paragraph in all_text_list:
                # remove chu thich
                if len(paragraph.split(' ')) > 15:
                    if paragraph[0].isnumeric() == True :
                        if paragraph[1].isupper() == True:
                            continue
                        elif paragraph[1].isnumeric() == True:
                            if paragraph[2].isupper() == True:
                                continue
                            elif paragraph[2].isnumeric() == True:
                                if paragraph[3].isupper() == True:
                                    continue
                #             
                if len(paragraph.split(' ')) > 25 :
                    all_text_list_process.append(paragraph)
                elif len(paragraph.split(' ')) > 15 :
                    if paragraph[0].isupper() == True and paragraph[-1] in stop_paragraph:
                        continue
                    elif paragraph[0].isupper() == True or paragraph[-1] in stop_paragraph or (paragraph[-1].isnumeric() == True and paragraph[-2] in stop_paragraph):
                        all_text_list_process.append(paragraph)
                    elif paragraph[0].isupper() == True or paragraph[-1] in stop_paragraph or (paragraph[-1].isnumeric() == True and paragraph[-2].isnumeric()== True and paragraph[-3] in stop_paragraph):
                        all_text_list_process.append(paragraph)
                    elif paragraph[0].isupper() == True or paragraph[-1] in stop_paragraph or (paragraph[-1].isnumeric() == True and paragraph[-2].isnumeric()== True and paragraph[-3].isnumeric()== True and paragraph[-4] in stop_paragraph):
                        all_text_list_process.append(paragraph)
                else:
                    continue
    for idx, paragraph in enumerate(all_text_list_process):
        try:
            if paragraph[0].isupper() == True and paragraph[-1] in stop_paragraph:
                all_text  = all_text+'\n'+ all_text_list_process[idx]
            elif paragraph[0].isupper() == True and paragraph[-1].isalpha() == False and paragraph[-2] in stop_paragraph:
                all_text  = all_text+'\n'+ all_text_list_process[idx]
            elif paragraph[0].isupper() == True and  all_text_list_process[idx+1][0].isupper() == False and paragraph[-1] not in stop_paragraph:
                    if len(all_text_list_process[idx+1])!=0 and all_text_list_process[idx+1][0].isupper() == False and all_text_list_process[idx+1][-1] in stop_paragraph:
                        all_text_list_process[idx] = all_text_list_process[idx]+' '+all_text_list_process[idx+1]
                        all_text  = all_text+'\n'+ all_text_list_process[idx]
                        all_text_list_process.pop(idx+1)
                    elif len(all_text_list_process[idx+1])!=0 and all_text_list_process[idx+1][0].isupper() == False and all_text_list_process[idx+1][-2] in stop_paragraph and [idx+1][-1].isnumeric() == True:
                        all_text_list_process[idx] = all_text_list_process[idx]+' '+all_text_list_process[idx+1]
                        all_text  = all_text+'\n'+ all_text_list_process[idx]
                        all_text_list_process.pop(idx+1)
                    elif len(all_text_list_process[idx+1])!=0 and all_text_list_process[idx+1][0].isupper() == False and all_text_list_process[idx+1][-1].isnumeric() == True and all_text_list_process[idx+1][-2].isnumeric()== True and all_text_list_process[idx+1][-3] in stop_paragraph:
                        all_text_list_process[idx] = all_text_list_process[idx]+' '+all_text_list_process[idx+1]
                        all_text  = all_text+'\n'+ all_text_list_process[idx]
                        all_text_list_process.pop(idx+1)
                    elif len(all_text_list_process[idx+1])!=0 and all_text_list_process[idx+1][0].isupper() == False and all_text_list_process[idx+1][-1].isnumeric() == True and all_text_list_process[idx+1][-2].isnumeric()== True and all_text_list_process[idx+1][-3].isnumeric()== True and all_text_list_process[idx+1][-4] in stop_paragraph:
                        all_text_list_process[idx] = all_text_list_process[idx]+' '+all_text_list_process[idx+1]
                        all_text  = all_text+'\n'+ all_text_list_process[idx]
                        all_text_list_process.pop(idx+1)
        except:
            continue
    all_text_unprocess = '\n'.join(all_text_list_unprocess)
    if len(all_text.split(' ')) / len(all_text_unprocess.split(' ')) <= 0.35:
        print('ok new process')
        print(str(len(all_text.split(' ')) / len(all_text_unprocess.split(' '))))
        temp_doc = tempfile.NamedTemporaryFile(suffix='.docx', delete=True)
        print("Created file for doc is:", temp_doc)
        print("Name of the file is:", temp_doc.name)
        doc = aw.Document(file_name)

        # Save the document to DOCX format.
        doc.save(temp_doc.name)
        # text = process_doc_text(temp_doc.name+'.docx')
        text = process_doc_text(temp_doc.name,page_from,page_to)
        temp_doc.close()
        return text
    print('ok old process')
    return all_text
def process_doc_text(file_name,page_from,page_to):
    doc = aw.Document(file_name)
    pageCount = doc.page_count
    #  truong hop page_to=page_from ==0 thi logic nay bi sai
    # dan den truong hop xet 1 trang se bi loi
    if (page_to == pageCount and page_from == 0)or page_to > pageCount:
        print('the same ext doc')
    else:
        temp_doc = tempfile.NamedTemporaryFile(suffix='.docx', delete=True)
        # print("Created file for doc is:", temp_doc)
        # print("Name of the file is:", temp_doc.name)
        if page_to == page_from and page_from != pageCount:
            extractedPages = doc.extract_pages(page_from, 1)
            extractedPages.save(temp_doc.name)
        elif  page_to == page_from and page_from == pageCount:
            extractedPages = doc.extract_pages(page_from-1, 1)
            extractedPages.save(temp_doc.name)          
        else:
            extractedPages = doc.extract_pages(page_from, (page_to-page_from+1))
            extractedPages.save(temp_doc.name)
        file_name = temp_doc.name
    doc = docx.Document(file_name)
    all_paras = doc.paragraphs
    all_text = ''
    for para in all_paras:
        text = re.sub(' +', ' ', para.text)
        text = text.replace('\t',"")
        # strencode = text.encode("ascii", "ignore")
        strencode = text.encode()
            #decode() method
        text = strencode.decode()
        if get_length(text) > 15:
            all_text = all_text +'\n'+text
    return all_text

# text = '''Anxiety affects quality of life in those living with Parkinson's disease (PD) more so than overall cognitive status, motor deficits, apathy, and depression [1–3]. Although anxiety and depression are often related and coexist in PD patients [4], recent research suggests that anxiety rather than depression is the most prominent and prevalent mood disorder in PD [5, 6]. Yet, our current understanding of anxiety and its impact on cognition in PD, as well as its neural basis and best treatment practices, remains meager and lags far behind that of depression [7].
# Overall, neuropsychiatric symptoms in PD have been shown to be negatively associated with cognitive performance. For example, higher depression scores have been correlated with lower scores on the Mini-Mental State Exam (MMSE) [8, 9] as well as tests of memory and executive functions (e.g., attention) [10–14]. Likewise, apathy and anhedonia in PD patients have been associated with executive dysfunction [10, 15–23]. However, few studies have specifically investigated the relationship between anxiety and cognition in PD.
# One study showed a strong negative relationship "between" MMSE anxiety (both state and trait) and overall cognitive performance (measured by the total of the repeatable battery for the assessment of neuropsychological status index) within a sample of 27 PD patients [24]. Furthermore, trait anxiety was negatively associated with each of the cognitive domains assessed by the RBANS (i.e., immediate memory, visuospatial construction, language, attention, and delayed memory). Two further studies have examined whether anxiety differentially affects cognition in patients with left-sided dominant PD (LPD) versus right-sided dominant PD (RPD); however, their findings were inconsistent. 
# The first study found that working memory performance was worse in LPD patients with anxiety compared to RPD patients with anxiety [25], whereas the second study reported that, in LPD, apathy but not anxiety was associated with performance on nonverbally mediated executive functions and visuospatial tasks (e.g., TMT-B, WMS-III spatial span), while in RPD, anxiety but not apathy significantly correlated with performance on verbally mediated tasks (e.g., clock reading test and Boston naming test) [15]. 
# Furthermore, anxiety was significantly correlated with neuropsychological measures of attention and executive and visuospatial functions [15]. Taken together, it is evident that there are limited and inconsistent findings describing the relationship between anxiety and cognition in PD and more specifically how anxiety might influence particular domains of cognition such as attention and memory and executive functioning. It is also striking that, to date, no study has examined the influence of anxiety on cognition in PD by directly comparing groups of PD patients with and without anxiety while excluding depression. This was the primary objective of the current study.'''      

# topics = [["MMSE","anxiety,Parkinson"],["executive,dysfunction","RBANS"]]

# get_raw_text_by_topic(topics,text)
